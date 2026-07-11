"""
导出功能 — 生成 Word (.docx) 和 PDF 文件。
"""
import os
import tempfile
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


EXPORT_DIR = os.path.join(os.path.dirname(__file__), "..", "exports")


def export_book(book_title: str, chapters: list[dict], format: str = "docx") -> str:
    """导出回忆录，返回文件路径。"""
    os.makedirs(EXPORT_DIR, exist_ok=True)

    if format == "docx":
        return _export_docx(book_title, chapters)
    else:
        return _export_pdf(book_title, chapters)


def _export_docx(book_title: str, chapters: list[dict]) -> str:
    """生成 Word 文档。"""
    doc = Document()

    # ── 设置默认字体 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ── 封面：书名 ──
    # 空几行
    for _ in range(6):
        doc.add_paragraph("")

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(book_title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4A, 0x37, 0x28)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 我的人生回忆录 ——")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()

    # ── 正文各章 ──
    for i, ch in enumerate(chapters):
        # 章节标题
        heading = doc.add_heading(ch["title"], level=1)
        for run in heading.runs:
            run.font.name = "SimHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.color.rgb = RGBColor(0x4A, 0x37, 0x28)

        # 章节正文
        content = ch.get("content", "")
        paragraphs = content.split("\n")
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(para_text.strip())
                run.font.size = Pt(12)
                run.font.name = "SimSun"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 每章结尾加分页（最后一章不加）
        if i < len(chapters) - 1:
            doc.add_page_break()

    # ── 保存 ──
    filepath = os.path.join(EXPORT_DIR, "memoir.docx")
    doc.save(filepath)
    return filepath


def _export_pdf(book_title: str, chapters: list[dict]) -> str:
    """生成 PDF 文件（用 weasyprint 把 HTML 转 PDF）。"""
    # 构建 HTML
    chapters_html = ""
    for ch in chapters:
        content = ch.get("content", "").replace("\n", "<br>")
        chapters_html += f"""
        <div class="chapter">
            <h2>{ch['title']}</h2>
            <div class="content">{content}</div>
        </div>
        """

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>
    @page {{
        size: A5;
        margin: 2cm 1.8cm;
    }}
    body {{
        font-family: "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei", "SimSun", sans-serif;
        font-size: 12pt;
        line-height: 1.8;
        color: #4A3728;
    }}
    .cover {{
        text-align: center;
        padding-top: 30%;
        page-break-after: always;
    }}
    .cover h1 {{
        font-size: 24pt;
        color: #4A3728;
    }}
    .cover .sub {{
        font-size: 13pt;
        color: #8B7355;
        margin-top: 1em;
    }}
    .chapter {{
        page-break-before: always;
    }}
    .chapter h2 {{
        font-size: 16pt;
        color: #C46B4A;
        margin-bottom: 1.5em;
        text-align: center;
    }}
    .chapter .content {{
        text-indent: 2em;
        text-align: justify;
    }}
</style>
</head>
<body>
<div class="cover">
    <h1>{book_title}</h1>
    <p class="sub">—— 我的人生回忆录 ——</p>
</div>
{chapters_html}
</body>
</html>"""

    # 用 weasyprint 渲染
    from weasyprint import HTML

    filepath = os.path.join(EXPORT_DIR, "memoir.pdf")
    HTML(string=html).write_pdf(filepath)
    return filepath
